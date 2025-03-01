from datetime import datetime, time
from msvcrt import getch
from sre_parse import parse_template
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login
from .forms import LoginForm, SignUpForm, ParticipantProfileForm,FileUploadForm, ScheduleSessionForm, ScheduledSessionForm,MentorAvailabilityForm,ChapterQuizForm, FinalQuizForm,JobSubmitForm, MentorProfileForm, MentorRatingForm, JobUploadForm, FeedbackForm, ResourceSearchForm, ResourceUploadForm, RatingForm, MentorshipRequestForm, SessionForm, ChatMessageForm, JobApplicationForm, MentorshipResponseForm, ThreadForm, ChangeEmailForm, PasswordChangeForm, DeleteAccountForm
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.admin.views.decorators import staff_member_required
from .models import User, MenteeActivity, Enrollment, Question, Message, ParticipantProfile, ScheduledSession, MentorAvailability, Rating, UserProgress, Chapter, Quiz, Certificate, MentorProfile, Feedback, Progress, Resource, MentorshipRequest, ChatMessage, Session, JobListing, SavedJob, JobApplication, Notification, Thread, Comment, Vote
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.http import JsonResponse
from django.db.models import Avg
from .google_calendar import get_google_calendar_service, create_calendar_event
from datetime import timedelta
from django.utils.timezone import now
from django.core.exceptions import PermissionDenied
from django.urls import reverse
from django.http import HttpResponseForbidden
from django.contrib.auth import update_session_auth_hash, logout
from django.contrib import messages
from django.core.mail import send_mail
from django.conf import settings
from reportlab.lib.pagesizes import letter, landscape
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from capstone.utils import generate_certificate
import os
from reportlab.lib import colors
from reportlab.lib.utils import simpleSplit
from reportlab.lib.colors import black, blue
from django.forms import inlineformset_factory
import logging
from io import BytesIO
from django.core.files import File
from django.http import FileResponse
from django.core.files.base import ContentFile
from reportlab.pdfbase.ttfonts import TTFont
import re
from reportlab.pdfbase import pdfmetrics
import openai
import json
import PyPDF2
import docx
from django.core.files.storage import default_storage
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
from .models import QuizAttempt
from .utils import extract_text_from_file 
from .utils import translate_text
from django.core.mail import EmailMessage
from .ai_utils import generate_quiz_questions
import urllib.parse
from django.utils.timezone import localtime, make_aware
from django.utils.html import strip_tags


logger = logging.getLogger(__name__)

def homepage(request):
    return render(request, 'capstone/homepage.html')

def about_us(request):
    return render(request, 'capstone/about_us.html')

def features(request):
    return render(request, 'capstone/features.html')

def contact_us(request):
    return render(request, 'capstone/contact_us.html')

def login_view(request):
    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            username = form.cleaned_data['username']
            password = form.cleaned_data['password']
            user_type = form.cleaned_data['user_type']
            
            
            user = authenticate(request, username=username, password=password)
            
            if user is not None and user.user_type == user_type:
                login(request, user)
                if user.user_type == 'participant':
                    return redirect('participant_dashboard')
                elif user.user_type == 'mentor':
                    return redirect('mentor_dashboard')
            else:
                form.add_error(None, 'Invalid email, password, or user type')
    else:
        form = LoginForm()
    # Provide a default language if user is not authenticated
    language_preference = request.user.language_preference 
    
    return render(request, 'capstone/login.html', {'form': form, 'language_preference': language_preference})


def signup_view(request):
    if request.method == 'POST':
        form = SignUpForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.set_password(form.cleaned_data['password'])
            user.username = form.cleaned_data['username']
            user.save()

            
            send_mail(
                'Welcome to Hustle Platform!',
                f'Hi {user.full_name},\n\nWelcome to Hustle Platform! 🚀\n\nExplore mentorship, learning resources, and job opportunities.\n\nHappy Learning,\nThe Hustle Platform Team',
                settings.DEFAULT_FROM_EMAIL,
                [user.email],
                fail_silently=False,
            )

            return redirect('login')
    else:
        form = SignUpForm()
    return render(request, 'capstone/signup.html', {'form': form})

def mentor_terms(request):
    return render(request, 'capstone/mentor_terms.html')

def participant_terms(request):
    return render(request, 'capstone/participant_terms.html')

@login_required
def participant_dashboard(request, participant_id=None):
    if participant_id:
        participant = get_object_or_404(User, id=participant_id, user_type='participant')
    else:
        participant = request.user  

    mentorship_request = MentorshipRequest.objects.filter(
        mentee=participant, status='approved'
    ).first()  

    mentor = mentorship_request.mentor if mentorship_request else None

    notifications = Notification.objects.filter(user=participant, is_seen=False).order_by('-created_at')

    context = {
        'user': participant,
        'mentor': mentor,  
        'notifications': notifications,
    }
    return render(request, 'capstone/participant_dashboard.html', context)


@login_required
def mentor_dashboard(request):
    if request.user.user_type != 'mentor' and not request.user.is_superuser:
        return redirect('login')   
    notifications = Notification.objects.filter(user=request.user, is_seen= False).order_by('-created_at')
    mentees = User.objects.filter(
        id__in=MentorshipRequest.objects.filter(mentor=request.user, status='approved').values_list('mentee_id', flat=True)
    )
    mentorship_requests = MentorshipRequest.objects.filter(mentor=request.user, status='pending')
    messages = Message.objects.filter(mentor=request.user).order_by('-timestamp')[:5]
    context = {
        'user': request.user,
        'notifications': notifications,
        'mentees': mentees,
        'messages': messages,
        'mentorship_requests': mentorship_requests,
    }

    return render(request, 'capstone/mentor_dashboard.html', context)

@login_required
def participant_profile(request, participant_id):
    participant = get_object_or_404(User, id=participant_id, user_type='participant')
    
    # Ensure we get the latest saved profile
    profile, created = ParticipantProfile.objects.get_or_create(user=participant)
    profile.refresh_from_db()  # 🔥 Ensure it fetches the latest changes

    # ✅ Fetch user's progress in courses
    user_progress = UserProgress.objects.filter(user=participant).select_related("resource")

    # Fetch completed certificates
    certificates = Certificate.objects.filter(user=participant)

    return render(request, 'capstone/participant_profile.html', {
        'participant': participant,
        'profile': profile,  # Ensure profile is passed to the template
        'certificates': certificates,
        'user_progress': user_progress,
    })

@login_required
def edit_participant_profile(request, participant_id):
    participant = get_object_or_404(User, id=participant_id, user_type='participant')
    
    # Fetch the latest profile instance
    profile, created = ParticipantProfile.objects.get_or_create(user=participant)

    if request.method == 'POST':
        form = ParticipantProfileForm(request.POST, request.FILES, instance=profile)
        if form.is_valid():
            form.save()
            profile.refresh_from_db()  # 🔥 Fetch latest saved data
            return redirect('participant_profile', participant_id=participant_id)
    else:
        form = ParticipantProfileForm(instance=profile)

    return render(request, 'capstone/edit_participant_profile.html', {
        'form': form,
        'participant': participant
    })

@login_required
def mentor_profile(request, mentor_id):
    mentor = get_object_or_404(User, id=mentor_id, user_type='mentor')
    
    # Ensure we get the latest saved profile
    profile, created = MentorProfile.objects.get_or_create(user=mentor)
    profile.refresh_from_db()  # 🔥 Ensure it fetches the latest changes

    # ✅ Fetch Mentor Availability
    availability_slots = MentorAvailability.objects.filter(mentor=mentor).order_by("date", "start_time")

    mentees = User.objects.filter(
        id__in=MentorshipRequest.objects.filter(mentor=mentor, status='approved').values_list('mentee_id', flat=True)
    )
    feedbacks = Feedback.objects.filter(mentor=mentor)

    return render(request, 'capstone/mentor_profile.html', {
        'profile': profile,
        'mentor': mentor,
        'mentees': mentees,
        'feedbacks': feedbacks,
        'availability_slots': availability_slots,
    })

@login_required
def edit_mentor_profile(request, mentor_id):
    mentor = get_object_or_404(User, id=mentor_id, user_type='mentor')

    # Fetch the latest profile instance
    profile, created = MentorProfile.objects.get_or_create(user=mentor)

    if request.method == 'POST':
        form = MentorProfileForm(request.POST, request.FILES, instance=profile)
        if form.is_valid():
            form.save()
            profile.refresh_from_db()  # 🔥 Fetch latest saved data
            return redirect('mentor_profile', mentor_id=mentor.id)
    else:
        form = MentorProfileForm(instance=profile)

    return render(request, 'capstone/edit_mentor_profile.html', {
        'form': form,
        'mentor': mentor
    })


@login_required
def submit_feedback(request, mentor_id):
    mentor = get_object_or_404(User, id=mentor_id)
    if request.method == 'POST':
        form = FeedbackForm(request.POST)
        if form.is_valid():
            feedback = form.save(commit=False)
            feedback.mentor = mentor
            feedback.mentee = request.user
            feedback.save()
            return redirect('participant_dashboard')
    else:
        form = FeedbackForm()
    return render(request, 'capstone/submit_feedback.html', {'form': form, 'mentor': mentor})

@login_required
def learning_resources(request):
    resources = Resource.objects.all()
    user_progress = UserProgress.objects.filter(user=request.user)
    
    return render(request, "capstone/learning_resources.html", {
        "resources": resources,
        "user_progress": user_progress
    })   

def mlearning_resources(request):
    base_template = "mentor-base.html" if request.user.is_authenticated and request.user.user_type == "mentor" else "participant-base.html"
    resources = Resource.objects.all().annotate(avg_rating=Avg('ratings__rating')).order_by('id')
    mentorship_requests = MentorshipRequest.objects.filter(mentee=request.user)
    form = ResourceSearchForm(request.GET)

    

    if form.is_valid():
        keyword = form.cleaned_data.get('keyword')
        topic = form.cleaned_data.get('topic')
        format = form.cleaned_data.get('format')
        difficulty = form.cleaned_data.get('difficulty')

        if keyword:
            resources = resources.filter(title__icontains=keyword)
        if topic:
            resources = resources.filter(topic__icontains=topic)
        if format:
            resources = resources.filter(format=format)
        if difficulty:
            resources = resources.filter(difficulty=difficulty)
    
    
    paginator = Paginator(resources, 5)  
    page_number = request.GET.get('page')
    if not page_number: 
        page_number = 1
    try:
        resources = paginator.page(page_number)
    except PageNotAnInteger:
        resources = paginator.page(1)
    except EmptyPage:
        resources = paginator.page(paginator.num_pages)

    resources = Resource.objects.all()

    context = {
        'resources': resources,
        'form': form,
        'mentorship_requests': mentorship_requests,
        "base_template": base_template,
    }
    return render(request, 'capstone/mentor-learning_resources.html', context)


@login_required
def rate_resource(request, resource_id):
    resource = get_object_or_404(Resource, id=resource_id)
    
    # Check if user has already rated
    user_rating = Rating.objects.filter(resource=resource, user=request.user).first()

    if request.method == 'POST':
        form = RatingForm(request.POST, instance=user_rating)  # Pre-fill if rating exists
        if form.is_valid():
            rating = form.save(commit=False)
            rating.resource = resource
            rating.user = request.user
            rating.save()

            print(f"⭐ Rating saved: {rating.rating}/5 by {request.user.full_name} for {resource.title}")
            return redirect('learning_resources')

        else:
            print("Form errors:", form.errors)  # Debugging
    else:
        form = RatingForm(instance=user_rating)

    return render(request, 'capstone/rate_resource.html', {
        'form': form,
        'resource': resource,
        'user_rating': user_rating
    })


@login_required
def submit_mentorship_request(request):
    mentors = User.objects.filter(user_type='mentor')  

    if request.method == 'POST':
        form = MentorshipRequestForm(request.POST)
        if form.is_valid():
            mentorship_request = form.save(commit=False)
            mentorship_request.mentee = request.user
            mentorship_request.save()

             # Send email notification to the mentor
            mentor_email = mentorship_request.mentor.email  # Ensure mentor field exists in MentorshipRequest model
            subject = "New Mentorship Request"
            message = f"Hello {mentorship_request.mentor.first_name},\n\n" \
                      f"You have received a new mentorship request from {request.user.first_name} {request.user.last_name}.\n\n" \
                      f"Please log in to your dashboard to respond.\n\n" \
                      f"Best regards,\nYour Mentorship Team"

            send_mail(
                subject,
                message,
                settings.DEFAULT_FROM_EMAIL,  # Ensure this is set in your settings.py
                [mentor_email],
                fail_silently=False,
            )
            return redirect('participant_dashboard')
    else:
        form = MentorshipRequestForm()

    return render(request, 'capstone/submit_mentorship_request.html', {
        'form': form,
        'mentors': mentors,  
    })

@login_required
def schedule_session(request, request_id):
    mentorship_request = get_object_or_404(MentorshipRequest, id=request_id, mentee=request.user)
    if request.method == 'POST':
        form = SessionForm(request.POST)
        if form.is_valid():
            session = form.save(commit=False)
            session.mentorship_request = mentorship_request
            session.save()

            
            event_link = create_calendar_event(
                summary=f"Session with {mentorship_request.mentor.full_name}",
                description=session.notes,
                start_time=session.scheduled_time.isoformat(),
                end_time=(session.scheduled_time + datetime.timedelta(hours=1)).isoformat(),
                timezone='UTC',
                attendees=[{'email': mentorship_request.mentor.email}]
            )

            
            Notification.objects.create(
                user=mentorship_request.mentor,
                message=f"A session has been scheduled with {request.user.full_name} on {session.scheduled_time}. Event link: {event_link}"
            )
            Notification.objects.create(
                user=request.user,
                message=f"A session has been scheduled with {mentorship_request.mentor.full_name} on {session.scheduled_time}. Event link: {event_link}"
            )

            # Send email notifications
            subject = "New Mentorship Session Scheduled"
            mentee_message = f"Hello {request.user.first_name},\n\n" \
                             f"Your mentorship session with {mentorship_request.mentor.full_name} is scheduled on {session.scheduled_time}.\n\n" \
                             f"You can join using this link: {event_link}\n\n" \
                             f"Best regards,\nYour Mentorship Team"

            mentor_message = f"Hello {mentorship_request.mentor.first_name},\n\n" \
                             f"You have a mentorship session scheduled with {request.user.full_name} on {session.scheduled_time}.\n\n" \
                             f"Join the session here: {event_link}\n\n" \
                             f"Best regards,\nYour Mentorship Team"

            send_mail(
                subject,
                mentor_message,
                settings.DEFAULT_FROM_EMAIL,
                [mentorship_request.mentor.email],
                fail_silently=False,
            )

            send_mail(
                subject,
                mentee_message,
                settings.DEFAULT_FROM_EMAIL,
                [request.user.email],
                fail_silently=False,
            )


            return redirect('participant_dashboard')
    else:
        form = SessionForm()
    return render(request, 'capstone/schedule_session.html', {'form': form, 'mentorship_request': mentorship_request})

def chat_with_mentor(request, request_id):
    mentorship_request = get_object_or_404(MentorshipRequest, id=request_id, mentee=request.user)
    chat_messages = ChatMessage.objects.filter(mentorship_request=mentorship_request).order_by('timestamp')

    if request.method == 'POST':
        form = ChatMessageForm(request.POST)
        if form.is_valid():
            chat_message = form.save(commit=False)
            chat_message.mentorship_request = mentorship_request
            chat_message.sender = request.user
            chat_message.save()

            
            if not mentorship_request.mentor.is_active:
                send_mail(
                    'New Message from Your Mentee',
                    f'Hi {mentorship_request.mentor.full_name},\n\nYou have a new message from {request.user.full_name}.\nLog in to check it.\n\nBest,\nHustle Platform Team',
                    settings.DEFAULT_FROM_EMAIL,
                    [mentorship_request.mentor.email],
                    fail_silently=False,
                )

            return redirect('chat_with_mentor', request_id=request_id)
    else:
        form = ChatMessageForm()

    return render(request, 'capstone/chat_with_mentor.html', {'form': form, 'chat_messages': chat_messages, 'mentorship_request': mentorship_request})


@login_required
def manage_mentorship_requests(request):
    mentorship_requests = MentorshipRequest.objects.filter(mentor=request.user, status='pending')
    return render(request, 'capstone/manage_mentorship_requests.html', {'mentorship_requests': mentorship_requests})

@login_required
def approve_mentorship_request(request, request_id):
    mentorship_request = get_object_or_404(MentorshipRequest, id=request_id, mentor=request.user)
    if request.method == 'POST':
        form = MentorshipResponseForm(request.POST, instance=mentorship_request)
        if form.is_valid():
            mentorship_request.status = 'approved'  
            form.save()

        
            mentee_profile, created = ParticipantProfile.objects.get_or_create(user=mentorship_request.mentee)
            mentee_profile.mentor = request.user  
            mentee_profile.save()

            
            Notification.objects.create(
                user=mentorship_request.mentee,
                message=f"Your mentorship request has been approved by {request.user.full_name}. Response: {mentorship_request.mentor_response}"
            )

            # Send email notification to the mentee
            subject = "Mentorship Request Approved"
            message = f"Hello {mentorship_request.mentee.first_name},\n\n" \
                      f"Your mentorship request has been approved by {request.user.full_name}.\n\n" \
                      f"Mentor's Response: {mentorship_request.mentor_response}\n\n" \
                      f"You can now communicate with your mentor and schedule sessions as needed.\n\n" \
                      f"Best regards,\nYour Mentorship Team"

            send_mail(
                subject,
                message,
                settings.DEFAULT_FROM_EMAIL,
                [mentorship_request.mentee.email],
                fail_silently=False,
            )


            messages.success(request, f"You are now mentoring {mentorship_request.mentee.full_name}.")
            return redirect('manage_mentorship_requests')
    else:
        form = MentorshipResponseForm(initial={'status': 'approved'})
    return render(request, 'capstone/mentorship_response.html', {'form': form, 'mentorship_request': mentorship_request})

@login_required
def decline_mentorship_request(request, request_id):
    mentorship_request = get_object_or_404(MentorshipRequest, id=request_id, mentor=request.user)
    if request.method == 'POST':
        form = MentorshipResponseForm(request.POST, instance=mentorship_request)
        if form.is_valid():
            mentorship_request.status = 'declined'  
            form.save()
            
            Notification.objects.create(
                user=mentorship_request.mentee,
                message=f"Your mentorship request has been declined by {request.user.full_name}. Response: {mentorship_request.mentor_response}"
            )

             # Send email notification to the mentee
            subject = "Mentorship Request Declined"
            message = f"Hello {mentorship_request.mentee.first_name},\n\n" \
                      f"We regret to inform you that your mentorship request has been declined by {request.user.full_name}.\n\n" \
                      f"Mentor's Response: {mentorship_request.mentor_response}\n\n" \
                      f"You may consider requesting another mentor or reapplying in the future.\n\n" \
                      f"Best regards,\nYour Mentorship Team"

            send_mail(
                subject,
                message,
                settings.DEFAULT_FROM_EMAIL,
                [mentorship_request.mentee.email],
                fail_silently=False,
            )
            return redirect('manage_mentorship_requests')
    else:
        form = MentorshipResponseForm(initial={'status': 'declined'})
    return render(request, 'capstone/mentorship_response.html', {'form': form, 'mentorship_request': mentorship_request})

def chat_with_mentee(request, request_id):
    mentorship_request = get_object_or_404(MentorshipRequest, id=request_id, mentor=request.user)
    chat_messages = ChatMessage.objects.filter(mentorship_request=mentorship_request).order_by('timestamp')

    if request.method == 'POST':
        form = ChatMessageForm(request.POST)
        if form.is_valid():
            chat_message = form.save(commit=False)
            chat_message.mentorship_request = mentorship_request
            chat_message.sender = request.user
            chat_message.save()

    
            if not mentorship_request.mentee.is_active:
                send_mail(
                    'New Message from Your Mentor',
                    f'Hi {mentorship_request.mentee.full_name},\n\nYou have a new message from {request.user.full_name}. Log in to check it.\n\nBest,\nHustle Platform Team',
                    settings.EMAIL_HOST_USER,
                    [mentorship_request.mentee.email],
                    fail_silently=False,
                )

            return redirect('chat_with_mentee', request_id=request_id)
    else:
        form = ChatMessageForm()

    return render(request, 'capstone/chat_with_mentee.html', {'form': form, 'chat_messages': chat_messages, 'mentorship_request': mentorship_request})

@login_required
def manage_schedule(request):
    sessions = Session.objects.filter(mentorship_request__mentor=request.user).order_by('scheduled_time')
    return render(request, 'capstone/manage_schedule.html', {'sessions': sessions})

def job_listings(request):
    job_listings = JobListing.objects.all()
    location = request.GET.get('location')
    industry = request.GET.get('industry')
    job_type = request.GET.get('job_type')

    if location:
        job_listings = job_listings.filter(location__icontains=location)
    if industry:
        job_listings = job_listings.filter(industry__icontains=industry)
    if job_type:
        job_listings = job_listings.filter(job_type=job_type)

        
    job_applications = JobApplication.objects.filter(user=request.user)



    context = {
        'job_listings': job_listings,
        'job_applications': job_applications, 
    }
    return render(request, 'capstone/job_listings.html', context)

@login_required
def save_job(request, job_id):
    job_listing = get_object_or_404(JobListing, id=job_id)
    SavedJob.objects.get_or_create(user=request.user, job_listing=job_listing)
    return redirect('job_listings')


@login_required
def apply_for_job(request, job_id):
    job_listing = get_object_or_404(JobListing, id=job_id)
    mentor = job_listing.employer  

    if request.method == 'POST':
        form = JobSubmitForm(request.POST, request.FILES) 
        if form.is_valid():
            job_application = form.save(commit=False)
            job_application.user = request.user
            job_application.job_listing = job_listing
            job_application.save()

            # ✅ Notify Mentor (Job Poster)
            Notification.objects.create(
                user=mentor,  
                message=f"📄 {request.user.full_name} has applied for your job listing: {job_listing.title}.",
                job_application=job_application,
            )

            # ✅ Send Email Notification to Mentor
            subject = f"📢 New Job Application - {job_listing.title}"
            email_body = f"""
                <p>Dear {mentor.full_name},</p>
                <p><strong>{request.user.full_name}</strong> has submitted an application for your job listing: <strong>{job_listing.title}</strong>.</p>
                <p>💼 Review applications by logging in to your dashboard.</p>
                <p><a href="https://your-platform.com/mentor-job-listings">🔍 View Applications</a></p>
                <br>
                <p>Best regards,<br> Hustle Platform Team</p>
            """

            email = EmailMessage(
                subject=subject,
                body=email_body,
                from_email=settings.DEFAULT_FROM_EMAIL,
                to=[mentor.email],
            )
            email.content_subtype = "html"  # Send as HTML email
            email.send(fail_silently=False)

            logger.info(f"✅ Job application submitted by {request.user.full_name} for {job_listing.title}")

            messages.success(request, "🎉 Job application submitted successfully!")
            return redirect('job_application_tracker')  
        else:
            logger.error(f"❌ Job application form errors: {form.errors}")
            messages.error(request, "❌ There was an error submitting your application. Please try again.")

    else:
        form = JobSubmitForm(initial={'job_listing': job_listing})

    return render(request, 'capstone/apply_for_job.html', {
        'form': form,
        'job_listing': job_listing,
    })


@login_required
def job_application_tracker(request):
    job_applications = JobApplication.objects.filter(user=request.user)
    context = {
        'job_applications': job_applications,
    }
    return render(request, 'capstone/job_application_tracker.html', context)



@login_required
def manage_applicants(request, job_id):
    job_listing = get_object_or_404(JobListing, id=job_id, employer=request.user)
    job_applications = JobApplication.objects.filter(job_listing=job_listing, status='submitted')
    context = {
        'job_listing': job_listing,
        'job_applications': job_applications,
    }
    return render(request, 'capstone/manage_applicants.html', context)


@login_required
def update_application_status(request, application_id, status):
    job_application = get_object_or_404(JobApplication, id=application_id, job_listing__employer=request.user)

    # ✅ Update application status
    job_application.status = status
    job_application.save()

    # ✅ Notify the participant
    Notification.objects.create(
        user=job_application.user,
        message=f"📢 Your application for <strong>{job_application.job_listing.title}</strong> has been <strong>{status.replace('_', ' ')}</strong>.",
    )

    # ✅ Send Email Notification to Participant
    subject = f"📢 Job Application Status Update - {job_application.job_listing.title}"
    email_body = f"""
        <p>Dear {job_application.user.full_name},</p>
        <p>Your job application for <strong>{job_application.job_listing.title}</strong> has been updated to: <strong>{status.replace('_', ' ')}</strong>.</p>
        <p>💼 You can check your application status by logging into your dashboard.</p>
        <p><a href="https://your-platform.com/job-application-tracker">🔍 View Application Status</a></p>
        <br>
        <p>Best regards,<br> Hustle Platform Team</p>
    """

    email = EmailMessage(
        subject=subject,
        body=email_body,
        from_email=settings.DEFAULT_FROM_EMAIL,
        to=[job_application.user.email],
    )
    email.content_subtype = "html"  # Send as HTML email

    try:
        email.send(fail_silently=False)
        logger.info(f"✅ Email sent successfully to {job_application.user.email} for job application update: {status}")
    except Exception as e:
        logger.error(f"❌ Email sending failed: {e}")
        messages.error(request, "Email notification could not be sent. Please check your email configuration.")

    messages.success(request, f"🎉 Application status updated to '{status.replace('_', ' ')}'. The applicant has been notified.")
    return JsonResponse({"message": f"Application status updated to {status}"})

@login_required
def mark_notification_seen(request, notification_id):
    notification = get_object_or_404(Notification, id=notification_id, user=request.user)
    notification.is_seen = True
    notification.save()
    return redirect('participant_dashboard')

@login_required
def upvote_message(request, message_id):
    message = get_object_or_404(ChatMessage, id=message_id)
    message.upvotes += 1
    message.save()
    return redirect('chat_with_mentor', request_id=message.mentorship_request.id)

@login_required
def downvote_message(request, message_id):
    message = get_object_or_404(ChatMessage, id=message_id)
    message.downvotes += 1
    message.save()
    return redirect('chat_with_mentor', request_id=message.mentorship_request.id)



@login_required
def google_callback(request):
    
    get_google_calendar_service()
    return redirect('participant_dashboard')

@login_required
def create_thread(request):
    if request.method == 'POST':
        form = ThreadForm(request.POST)
        if form.is_valid():
            thread = form.save(commit=False)
            thread.created_by = request.user  
            thread.save()

            # ✅ Notify all users
            notify_users(f"🆕 New discussion thread: {thread.title}", thread.id)

            return redirect('thread_detail', thread_id=thread.id)  
    else:
        form = ThreadForm()  
    return render(request, 'capstone/create_thread.html', {'form': form})

@login_required
def mcreate_thread(request):
    if request.method == 'POST':
        form = ThreadForm(request.POST)
        if form.is_valid():
            thread = form.save(commit=False)
            thread.created_by = request.user  
            thread.save()

            notify_users(f"🚀 A new discussion has started: '{thread.title}'! Join the conversation now! 💬", thread.id)


            return redirect('mentorthread_detail', thread_id=thread.id)  
    else:
        form = ThreadForm()  
    return render(request, 'capstone/mentor-create_thread.html', {'form': form})

@login_required
def reply_to_thread(request, thread_id):
    thread = Thread.objects.get(id=thread_id)
    if request.method == 'POST':
        text = request.POST.get('text')
        Comment.objects.create(thread=thread, user=request.user, text=text)

        # ✅ Notify all users
        notify_users(f"💬 Someone just joined the discussion in '{thread.title}'! Check it out. 👀", thread.id)


        return redirect('thread_detail', thread_id=thread.id)
    return render(request, 'capstone/reply_to_thread.html', {'thread': thread})

@login_required
def vote_comment(request, comment_id, vote_type):
    comment = get_object_or_404(Comment, id=comment_id)  

    
    existing_vote = Vote.objects.filter(user=request.user, comment=comment).first()
    if existing_vote:
        if existing_vote.vote_type != vote_type:
            if vote_type == 'upvote':
                comment.upvotes += 1
                comment.downvotes -= 1
            else:
                comment.upvotes -= 1
                comment.downvotes += 1
            existing_vote.vote_type = vote_type
            existing_vote.save()
    else:
        if vote_type == 'upvote':
            comment.upvotes += 1
        else:
            comment.downvotes += 1
        Vote.objects.create(user=request.user, comment=comment, vote_type=vote_type)

    
    comment.save()
    # ✅ Notify all users
    
    notify_users(
        f"🔥 The comment thread '{comment.thread.title}' received a {vote_type}! 🎉",)

    
    return redirect('thread_detail', thread_id=comment.thread.id)

@login_required
def mvote_comment(request, comment_id, vote_type):
    comment = get_object_or_404(Comment, id=comment_id)  

    
    existing_vote = Vote.objects.filter(user=request.user, comment=comment).first()
    if existing_vote:
        
        if existing_vote.vote_type != vote_type:
            
            if vote_type == 'upvote':
                comment.upvotes += 1
                comment.downvotes -= 1
            else:
                comment.upvotes -= 1
                comment.downvotes += 1
            existing_vote.vote_type = vote_type
            existing_vote.save()
    else:
        
        if vote_type == 'upvote':
            comment.upvotes += 1
        else:
            comment.downvotes += 1
        Vote.objects.create(user=request.user, comment=comment, vote_type=vote_type)

        notify_users(
        f"🔥 The comment thread '{comment.thread.title}' received a {vote_type}! 🎉",
        comment.id
    )

    
    comment.save()

    
    return redirect('mentorthread_detail', thread_id=comment.thread.id)


@login_required
def delete_comment(request, comment_id):
    comment = get_object_or_404(Comment, id=comment_id)
    user = request.user  

  
    is_mentor = hasattr(user, 'mentor')  
    is_admin = user.is_staff or user.is_superuser

    if user == comment.user or is_mentor or is_admin:
        thread_id = comment.thread.id
        comment.delete()
        return redirect(reverse('thread_detail', args=[thread_id]))
    
    
    return HttpResponseForbidden("You are not allowed to delete this comment.")

@login_required
def mdelete_comment(request, comment_id):
    comment = get_object_or_404(Comment, id=comment_id)
    user = request.user  

    
    is_mentor = hasattr(user, 'mentor')  
    is_admin = user.is_staff or user.is_superuser

    if user == comment.user or is_mentor or is_admin:
        thread_id = comment.thread.id
        comment.delete()
        return redirect(reverse('mentorthread_detail', args=[thread_id]))
    
    
    return HttpResponseForbidden("You are not allowed to delete this comment.")


def list_threads(request):
    
    threads = Thread.objects.all().order_by('-created_at')  
    return render(request, 'capstone/list_threads.html', {'threads': threads})

def thread_detail(request, thread_id):
    thread = get_object_or_404(Thread, id=thread_id)
    return render(request, 'capstone/thread_detail.html', {'thread': thread})

def vote_thread(request, thread_id, action):
    thread = get_object_or_404(Thread, id=thread_id)
    
    if action == "upvote":
        thread.upvotes += 1
    elif action == "downvote":
        thread.downvotes += 1
    
    notify_users(
        f"🔥The discussion thread '{thread.title}' received a {action}! 🎉",
        thread.id
    )
    thread.save()
    return redirect('list_threads') 

def mentorlist_threads(request):
    
    threads = Thread.objects.all().order_by('-created_at')  
    return render(request, 'capstone/mentor-list_threads.html', {'threads': threads})

def thread_detail(request, thread_id):
    thread = get_object_or_404(Thread, id=thread_id)
    return render(request, 'capstone/thread_detail.html', {'thread': thread})

def mentorthread_detail(request, thread_id):
    thread = get_object_or_404(Thread, id=thread_id)
    return render(request, 'capstone/mentor-thread_detail.html', {'thread': thread})

def vote_thread(request, thread_id, action):
    thread = get_object_or_404(Thread, id=thread_id)
    
    if action == "upvote":
        thread.upvotes += 1
    elif action == "downvote":
        thread.downvotes += 1
    
    thread.save()
    # Notify thread creator about the vote
    notify_users(
        f"🔥 The discussion thread '{thread.title}' received a {action}! 🎉",
        thread.id
    )
    
    return redirect('list_threads') 

def mvote_thread(request, thread_id, action):
    thread = get_object_or_404(Thread, id=thread_id)
    
    if action == "upvote":
        thread.upvotes += 1
    elif action == "downvote":
        thread.downvotes += 1
    
    thread.save()

    # Notify thread creator about the vote
    notify_users(
        f"🔥 The discussion thread '{thread.title}' received a {action}! 🎉",
        thread.id
    )
    return redirect('mentorlist_threads') 

@login_required
def delete_thread(request, id):
    thread = get_object_or_404(Thread, id=id)
    
    if request.user == thread.created_by or request.user.is_staff:
        thread.delete()

    return redirect('list_threads')

@login_required
def mdelete_thread(request, id):
    thread = get_object_or_404(Thread, id=id)
    
    if request.user == thread.created_by or request.user.is_staff:
        thread.delete()

    return redirect('mentorlist_threads')

@login_required
def account_settings(request):
    return render(request, 'capstone/account_settings.html')

@login_required
def maccount_settings(request):
    return render(request, 'capstone/mentor-account_settings.html')

@login_required
def change_email(request):
    if request.method == 'POST':
        form = ChangeEmailForm(request.POST)
        if form.is_valid():
            new_email = form.cleaned_data['new_email']
            request.user.email = new_email
            request.user.save()
            messages.success(request, 'Your email has been updated successfully.')
            return redirect('account_settings')
    else:
        form = ChangeEmailForm()
    return render(request, 'capstone/change_email.html', {'form': form})

@login_required
def mchange_email(request):
    if request.method == 'POST':
        form = ChangeEmailForm(request.POST)
        if form.is_valid():
            new_email = form.cleaned_data['new_email']
            request.user.email = new_email
            request.user.save()
            messages.success(request, 'Your email has been updated successfully.')
            return redirect('mentor-account_settings')
    else:
        form = ChangeEmailForm()
    return render(request, 'capstone/mentor-change_email.html', {'form': form})

@login_required
def change_password(request):
    if request.method == 'POST':
        form = PasswordChangeForm(request.user, request.POST)
        if form.is_valid():
            user = form.save()  
            update_session_auth_hash(request, user)  
            messages.success(request, 'Your password has been updated successfully.')
            return redirect('login')  
        else:
            messages.error(request, 'Please correct the error below.')
    else:
        form = PasswordChangeForm(request.user)
    return render(request, 'capstone/change_password.html', {'form': form})

@login_required
def mchange_password(request):
    if request.method == 'POST':
        form = PasswordChangeForm(request.user, request.POST)
        if form.is_valid():
            user = form.save()  
            update_session_auth_hash(request, user)  
            messages.success(request, 'Your password has been updated successfully.')
            return redirect('login')  
        else:
            messages.error(request, 'Please correct the error below.')
    else:
        form = PasswordChangeForm(request.user)
    return render(request, 'capstone/mentor-change_password.html', {'form': form})

@login_required
def delete_account(request):
    if request.method == 'POST':
        form = DeleteAccountForm(request.POST)
        if form.is_valid():
            request.user.delete()  
            logout(request)  
            messages.success(request, 'Your account has been deleted successfully.')
            return redirect('login')
    else:
        form = DeleteAccountForm()
    return render(request, 'capstone/delete_account.html', {'form': form})

@login_required
def mdelete_account(request):
    if request.method == 'POST':
        form = DeleteAccountForm(request.POST)
        if form.is_valid():
            request.user.delete()  
            logout(request)  
            messages.success(request, 'Your account has been deleted successfully.')
            return redirect('login')
    else:
        form = DeleteAccountForm()
    return render(request, 'capstone/mentor-delete_account.html', {'form': form})

from django.core.mail import send_mail
from django.contrib.auth.decorators import login_required, user_passes_test
from django.shortcuts import render, redirect
from django.contrib import messages
from django.conf import settings
from .models import User, Notification, JobListing
from .forms import JobUploadForm

@login_required
@user_passes_test(lambda u: u.user_type == 'mentor' or u.is_superuser)
def mentor_job_listings(request):
    """Mentors can upload and view job listings."""
    job_listings = JobListing.objects.filter(employer=request.user)

    if request.method == 'POST':
        form = JobUploadForm(request.POST)
        if form.is_valid():
            job = form.save(commit=False)
            job.employer = request.user
            job.save()

            # ✅ Fetch Participants & Their Emails
            participants = User.objects.filter(user_type='participant').exclude(email__isnull=True).exclude(email='')

            # Convert QuerySet to List (Fixes Email Issues)
            participant_emails = list(participants.values_list('email', flat=True))

            # 🔍 Debugging: Print Fetched Emails
            print("📧 Participant Emails:", participant_emails)

            if not participant_emails:
                print("🚨 No valid participant emails found! Skipping email sending.")
            else:
                # ✅ Send System Notification to Participants
                for participant in participants:
                    Notification.objects.create(
                        user=participant,
                        message=f"A new job has been posted: {job.title}."
                    )

                # ✅ Send Email Notification to Participants
                subject = "New Job Opportunity Available!"
                message = f"Hello,\n\nA new job opportunity titled '{job.title}' has been posted by {request.user.full_name}.\n\n" \
                          f"Log in to your dashboard to view and apply.\n\n" \
                          f"Best regards,\nYour Mentorship Team"

                try:
                    send_mail(
                        subject,
                        message,
                        settings.DEFAULT_FROM_EMAIL,
                        participant_emails,
                        fail_silently=False,
                    )
                    print("✅ Emails sent successfully!")
                except Exception as e:
                    print(f"🚨 Email Sending Failed: {e}")

            messages.success(request, "Job listing created successfully!")
            return redirect('mentor_job_listings')

    else:
        form = JobUploadForm()

    context = {
        'job_listings': job_listings,
        'form': form,
    }
    return render(request, 'capstone/mentor-job_listings.html', context)


@login_required
def rate_mentor(request, mentor_id):
    mentor = get_object_or_404(User, id=mentor_id, user_type='mentor')

    if request.method == 'POST':
        form = MentorRatingForm(request.POST)
        if form.is_valid():
            feedback = form.save(commit=False)
            feedback.mentor = mentor
            feedback.mentee = request.user  
            feedback.save()

            # Send email notification to the mentor
            subject = "New Mentor Rating Received"
            message = f"Hello {mentor.first_name},\n\n" \
                      f"You have received a new rating and feedback from {request.user.full_name}.\n\n" \
                      f"Please log in to your dashboard to view the feedback.\n\n" \
                      f"Best regards,\nYour Mentorship Team"

            send_mail(
                subject,
                message,
                settings.DEFAULT_FROM_EMAIL,
                [mentor.email],
                fail_silently=False,
            )

            messages.success(request, "Your feedback has been submitted successfully!")
            return redirect('participant_dashboard')

    else:
        form = MentorRatingForm()

    return render(request, 'capstone/rate_mentor.html', {'form': form, 'mentor': mentor})

@login_required
def notification_seen(request, notification_id, job_id):
    """ Marks a notification as seen and redirects the mentor to manage applicants. """
    notification = get_object_or_404(Notification, id=notification_id, user=request.user)
    notification.is_seen = True
    notification.save()

    return redirect('manage_applicants', job_id=job_id)

@login_required
def mark_notification_seen(request, notification_id):
    notification = get_object_or_404(Notification, id=notification_id, user=request.user)
    notification.is_seen = True
    notification.save()
    return redirect('mentor_dashboard')

@login_required
def pmark_notification_seen(request, notification_id):
    notification = get_object_or_404(Notification, id=notification_id, user=request.user)
    notification.is_seen = True
    notification.save()
    return redirect('participant_dashboard')

@login_required
def update_application_status(request, application_id, status):
    job_application = get_object_or_404(JobApplication, id=application_id, job_listing__employer=request.user)

    # ✅ Update application status
    # ✅ If the mentor accepts the application, store their email
    if status == "accepted":
        job_application.accepted_by_mentor_email = request.user.email  # Mentor's email
    else:
        job_application.accepted_by_mentor_email = None
    job_application.status = status
    job_application.save()

    # ✅ Notify the participant
    Notification.objects.create(
        user=job_application.user,  # Notify the applicant
        message=f"📢 Your application for '{job_application.job_listing.title}' has been {status.replace('_', ' ')}.",
    )

    # ✅ Send Email Notification to Participant
    subject = f"📢 Job Application Status Update - {job_application.job_listing.title}"
    email_body = f"""
        <p>Dear {job_application.user.full_name},</p>
        <p>Your job application for <strong>{job_application.job_listing.title}</strong> has been updated to: <strong>{status.replace('_', ' ')}</strong>.</p>
        <p>💼 You can check your application status by logging into your dashboard.</p>
        <p><a href="https://your-platform.com/job-application-tracker">🔍 View Application Status</a></p>
        <br>
        <p>Best regards,<br> Hustle Platform Team</p>
    """

    email = EmailMessage(
        subject=subject,
        body=email_body,
        from_email=settings.DEFAULT_FROM_EMAIL,
        to=[job_application.user.email],
    )
    email.content_subtype = "html"  # Send as HTML email

    try:
        email.send(fail_silently=False)
        logger.info(f"✅ Email sent successfully to {job_application.user.email} for job application update: {status}")
    except Exception as e:
        logger.error(f"❌ Email sending failed: {e}")
        messages.error(request, "Email notification could not be sent. Please check your email configuration.")

    # ✅ Remove application from Manage Applicants Page (Redirect back)
    messages.success(request, f"🎉 Application status updated to '{status.replace('_', ' ')}'. The applicant has been notified.")
    return redirect('manage_applicants', job_id=job_application.job_listing.id)


@login_required
def delete_job_listing(request, job_id):
    job_listing = get_object_or_404(JobListing, id=job_id, employer=request.user)

    if request.method == "POST":
        job_listing.delete()
        messages.success(request, "Job listing deleted successfully.")
        return redirect('mentor_job_listings')  # ✅ Redirect back to the job listings page

    return render(request, 'capstone/delete_job_listing.html', {'job_listing': job_listing})

@login_required
def delete_learning_resource(request, resource_id):
    resource = get_object_or_404(Resource, id=resource_id)

    if request.method == "POST":
        resource.delete()
        messages.success(request, "Resource deleted successfully.")
        return redirect('mlearning_resources')  # ✅ Redirect back to mentor's resource page

    return render(request, 'capstone/delete_learning_resource.html', {'resource': resource})

@login_required
def delete_certificate(request, certificate_id):
    certificate = get_object_or_404(Certificate, id=certificate_id, user=request.user)
    
    if certificate.certificate_file:
        certificate.certificate_file.delete()  # Delete file from storage

    certificate.delete()
    messages.success(request, "📜 Certificate deleted successfully.")
    return redirect('participant_profile', participant_id=request.user.id)


@login_required
def issue_certificate(request, resource_id):
    resource = get_object_or_404(Resource, id=resource_id)
    progress = UserProgress.objects.get(user=request.user, resource=resource)

    if progress.progress_percentage() < 100 or not progress.final_score:
        messages.error(request, "You must complete all chapters and pass the final quiz before getting a certificate.")
        return redirect("course_detail", resource_id=resource.id)
    
    certificate = generate_certificate(resource, progress.final_score)


    messages.success(request, "Certificate generated successfully! You will receive an email with further details.")
    return redirect("participant_profile", participant_id=request.user.id)


@login_required
def enroll_course(request, resource_id):
    resource = get_object_or_404(Resource, id=resource_id)
    progress, created = UserProgress.objects.get_or_create(user=request.user, resource=resource)
    return redirect("course_detail", resource_id=resource.id)


@login_required
def menroll_course(request, resource_id):
    resource = get_object_or_404(Resource, id=resource_id)
    progress, created = UserProgress.objects.get_or_create(user=request.user, resource=resource)
    return redirect("mcourse_detail", resource_id=resource.id)

@login_required
def course_detail(request, resource_id):
    """Displays course details with chapters, quizzes, progress, and final quiz eligibility."""
    resource = get_object_or_404(Resource, id=resource_id)
    progress, created = UserProgress.objects.get_or_create(user=request.user, resource=resource)
    user_lang = request.user.language_preference  # Get user's language
    translated_title = resource.translated_title.get(user_lang, resource.title) if isinstance(resource.translated_title, dict) else resource.title
    translated_description = resource.translated_description.get(user_lang, resource.description) if isinstance(resource.translated_description, dict) else resource.description

    # Check if the final quiz exists
    translated_chapters = [
        {
            "title": translate_text(chapter.title, user_lang),
            "content": translate_text(chapter.content, user_lang),
        }
        for chapter in resource.chapters.all()
    ]

    final_quiz = Quiz.objects.filter(resource=resource, is_final_quiz=True).first()
    final_quiz_exists = final_quiz is not None  # Boolean flag for template

    return render(request, "capstone/course_detail.html", {
        "resource": resource,
        "progress": progress,
        "final_quiz_exists": final_quiz_exists,
        "final_quiz": final_quiz,  # Pass the final quiz object
        "translated_title": translated_title,
        "translated_description": translated_description,
        "translated_chapters": translated_chapters,
    })

@login_required
def mcourse_detail(request, resource_id):
    """Displays course details with chapters, quizzes, progress, and final quiz eligibility."""
    resource = get_object_or_404(Resource, id=resource_id)
    progress, created = UserProgress.objects.get_or_create(user=request.user, resource=resource)

    # Check if the final quiz exists
    final_quiz = Quiz.objects.filter(resource=resource, is_final_quiz=True).first()
    final_quiz_exists = final_quiz is not None  # Boolean flag for template

    return render(request, "capstone/mcourse_detail.html", {
        "resource": resource,
        "progress": progress,
        "final_quiz_exists": final_quiz_exists,
        "final_quiz": final_quiz  # Pass the final quiz object
    })


@login_required
def mark_chapter_done(request, chapter_id):
    chapter = get_object_or_404(Chapter, id=chapter_id)
    progress = UserProgress.objects.get(user=request.user, resource=chapter.resource)

    if chapter not in progress.completed_chapters.all():
        progress.completed_chapters.add(chapter)
        progress.save()
    
    return redirect("course_detail", resource_id=chapter.resource.id)


def extract_chapters(file):
    """
    Extracts text from a PDF or DOCX file while keeping full chapter titles and content.
    """
    file_path = default_storage.save(f"temp/{file.name}", ContentFile(file.read()))
    file_path = default_storage.path(file_path)
    
    text = ""

    # Extract text based on file type
    if file.name.endswith(".pdf"):
        text = extract_text_from_pdf(file_path)
    elif file.name.endswith(".docx"):
        text = extract_text_from_docx(file_path)

    # Extract chapters correctly
    chapters = split_into_chapters_by_title(text)

    # Clean up temporary file
    default_storage.delete(file_path)

    return chapters

def extract_text_from_pdf(pdf_path):
    """Extracts text from a PDF file."""
    text = ""
    with open(pdf_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text.strip()

def extract_text_from_docx(docx_path):
    """Extracts text from a DOCX file."""
    doc = docx.Document(docx_path)
    return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

def split_into_chapters_by_title(text):
    """
    Splits text into chapters based on 'Chapter X: Title' format.
    Ensures that chapter titles are retained exactly as they are in the document.
    """
    # Regular expression pattern to detect "Chapter X: Title"
    chapter_pattern = re.compile(r"(Chapter\s+\d+:\s+[^\n]+)", re.IGNORECASE)

    # Find all matches for chapter titles
    matches = list(chapter_pattern.finditer(text))

    chapters = []

    for i in range(len(matches)):
        start = matches[i].start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        
        # Extract the full chapter text (title + content)
        chapter_text = text[start:end].strip()

        # Ensure the extracted text includes the full title only once
        if chapter_text:
            chapters.append(chapter_text)

    return chapters

    

@login_required
@user_passes_test(lambda u: u.user_type == 'mentor' or u.is_superuser)
def upload_resource(request):
    """
    Allows mentors to upload a course along with a file, automatically extracting 5 chapters.
    """
    if request.method == 'POST':
        form = ResourceUploadForm(request.POST, request.FILES)

        if form.is_valid():
            resource = form.save(commit=False)
            resource.created_by = request.user
            
            uploaded_file = request.FILES.get('resource_file')
            if uploaded_file:
                resource.resource_file = uploaded_file
                resource.save()

                # Extract chapters
                chapters_content = extract_chapters(resource.resource_file)

                # Store extracted chapters properly
                for i, chapter_text in enumerate(chapters_content, start=1):
                    title, content = chapter_text.split("\n", 1) if "\n" in chapter_text else (chapter_text, "")
                    
                    Chapter.objects.create(
                        resource=resource,
                        title=title.strip(),  # Store full chapter title
                        content=content.strip(),  # Store chapter content
                        chapter_number=i
                    )

            # ✅ Notify All Participants
            participants = User.objects.filter(user_type='participant')
            participant_emails = list(participants.values_list('email', flat=True))

            for participant in participants:
                Notification.objects.create(
                    user=participant,
                    message=f"A new course '{resource.title}' has been uploaded!",
                    is_seen=False
                )

                participant_emails.append(participant.email)

            # ✅ Send Email Notification to Participants
            if participant_emails:
                subject = "New Course Uploaded!"
                message = f"Hello,\n\nA new course titled '{resource.title}' has been uploaded by {request.user.full_name}.\n\n" \
                          f"Log in to your dashboard to explore the content.\n\n" \
                          f"Best regards,\nYour Mentorship Team"
                
                send_mail(
                    subject,
                    message,
                    settings.DEFAULT_FROM_EMAIL,
                    participant_emails,
                    fail_silently=False,
                )


            messages.success(request, "Course uploaded with 5 chapters successfully!")
            return redirect('mentor_dashboard')

        else:
            messages.error(request, "Error uploading resource. Please check the form.")

    else:
        form = ResourceUploadForm()

    return render(request, 'capstone/upload_resource.html', {'form': form})


@login_required
def take_quiz(request, quiz_id):
    """Handles quiz attempt, ensuring all five questions per quiz are displayed and scored correctly."""
    quiz = get_object_or_404(Quiz, id=quiz_id)
    progress, created = UserProgress.objects.get_or_create(user=request.user, resource=quiz.chapter.resource)

    previous_attempt = progress.quiz_attempts.get(str(quiz_id), None)  

    if request.method == 'POST':
        user_answers = {}
        correct_answers = 0
        feedback = {}

        for i, question in enumerate(quiz.questions, start=1):
            question_text = question.get("question")
            correct_answer = question.get("correct_answer")
            selected_answer = request.POST.get(f'answer_{i}')

            if selected_answer == correct_answer:
                correct_answers += 1

            feedback[question_text] = {
                "selected": selected_answer,
                "correct": correct_answer,
                "is_correct": selected_answer == correct_answer
            }

            user_answers[question_text] = selected_answer

        score_percentage = (correct_answers / len(quiz.questions)) * 100  

        # Store user attempt with feedback
        progress.quiz_attempts[str(quiz_id)] = {
            "answers": user_answers,
            "score": score_percentage,
            "feedback": feedback
        }
        progress.save()

        required_score = 4  

        if correct_answers >= required_score:
            progress.completed_quizzes.add(quiz)
            progress.save()
            messages.success(request, "🎉 Quiz passed successfully!")
        else:
            progress.retry_after = now() + timedelta(days=2)
            progress.save()
            messages.error(request, "❌ Quiz failed! Try again in 48 hours.")

        return redirect('course_detail', resource_id=quiz.chapter.resource.id)

    return render(request, 'capstone/take_quiz.html', {'quiz': quiz, 'previous_attempt': previous_attempt})

@login_required
def take_final_quiz(request, quiz_id):
    """Handles the final quiz attempt."""
    quiz = get_object_or_404(Quiz, id=quiz_id, is_final_quiz=True)
    progress, created = UserProgress.objects.get_or_create(user=request.user, resource=quiz.resource)

    previous_attempt = progress.quiz_attempts.get(str(quiz_id), None)

    if request.method == 'POST':
        user_answers = {}
        correct_answers = 0
        feedback = {}

        for i, question in enumerate(quiz.questions, start=1):
            question_text = question.get("question")
            correct_answer = question.get("correct_answer")
            selected_answer = request.POST.get(f'answer_{i}')

            # Check if answer is correct
            if selected_answer == correct_answer:
                correct_answers += 1

            # Store feedback per question
            feedback[question_text] = {
                "selected": selected_answer,
                "correct": correct_answer,
                "is_correct": selected_answer == correct_answer
            }

            user_answers[question_text] = selected_answer

        score_percentage = (correct_answers / len(quiz.questions)) * 100  

        # Store attempt with feedback
        progress.quiz_attempts[str(quiz_id)] = {
            "answers": user_answers,
            "score": score_percentage,
            "feedback": feedback  # Store feedback in the database
        }
        progress.final_score = score_percentage
        progress.save()

        required_score = 8  # Final quiz requires 8 correct answers

        if correct_answers >= required_score:
            progress.completed_quizzes.add(quiz)
            progress.save()
            messages.success(request, "🎉 Final Quiz passed successfully!")
        else:
            progress.retry_after = now() + timedelta(days=2)
            progress.save()
            messages.error(request, "❌ Final Quiz failed! Try again in 48 hours.")

        return redirect('course_detail', resource_id=quiz.resource.id)

    return render(request, 'capstone/take_final_quiz.html', {'quiz': quiz, 'previous_attempt': previous_attempt})

@login_required
def check_progress(request, resource_id):
    resource = get_object_or_404(Resource, id=resource_id)
    progress, created = UserProgress.objects.get_or_create(user=request.user, resource=resource)

    total_chapters = resource.chapters.count()
    completed_chapters = progress.completed_chapters.count()
    total_quizzes = Quiz.objects.filter(resource=resource).count()
    completed_quizzes = progress.completed_quizzes.count()

    chapter_percentage = (completed_chapters / total_chapters) * 50 if total_chapters else 0
    quiz_percentage = (completed_quizzes / total_quizzes) * 50 if total_quizzes else 0
    total_progress = chapter_percentage + quiz_percentage

    if total_progress >= 100:
        certificate, created = Certificate.objects.get_or_create(
            user=request.user, resource=resource,
            defaults={"issued_at": now()}
        )

        if created or not certificate.certificate_file:
            buffer = BytesIO()
            c = canvas.Canvas(buffer, pagesize=landscape(letter))  # **Set to Landscape**
            width, height = landscape(letter)  # Get width & height for centering

            c.setTitle("Certificate of Completion")

            # **Certificate Design (Now Adjusted for Landscape)**
            c.setFont("Helvetica-Bold", 36)
            c.drawCentredString(width / 2, height - 100, "Certificate of Completion")

            c.setFont("Helvetica", 18)
            c.drawCentredString(width / 2, height - 160, "This is proudly presented to")

            c.setFont("Helvetica-Bold", 28)
            c.setFillColor("blue")
            c.drawCentredString(width / 2, height - 210, f"{request.user.get_full_name()}")
            c.setFillColor("black")

            c.setFont("Helvetica", 16)
            c.drawCentredString(width / 2, height - 260, "for successfully completing the course:")

            c.setFont("Helvetica-Bold", 20)
            c.drawCentredString(width / 2, height - 300, f"{resource.title}")

            c.setFont("Helvetica", 14)
            c.drawCentredString(width / 2, height - 350, "Issued by: Hustle Platform")

            c.setFont("Helvetica", 14)
            c.drawCentredString(width / 2, height - 370, f"Date: {now().strftime('%Y-%m-%d')}")

            c.showPage()
            c.save()

            buffer.seek(0)
            certificate_filename = f"certificates/{request.user.username}_{resource.id}_certificate.pdf"
            certificate.certificate_file.save(certificate_filename, ContentFile(buffer.read()))
            certificate.save()

            # ✅ **Send Email Notification with Certificate Link**
        email_subject = "🎉 Congratulations! Your Certificate is Ready!"
        email_body = f"""
        <p>Dear {request.user.get_full_name()},</p>
        <p>🎉 Congratulations! You have successfully completed the course <strong>{resource.title}</strong>.</p>
        <p>Your certificate is now available for download. Click go and collect it in your profile</p>
        <p>Thank you for learning with Hustle Platform! Keep striving for excellence. 🚀</p>
        <p>Best regards,<br><strong>Hustle Platform Team</strong></p>
        """

        email = EmailMessage(
            subject=email_subject,
            body=email_body,
            from_email=settings.DEFAULT_FROM_EMAIL,
            to=[request.user.email],
        )
        email.content_subtype = "html"
        email.send(fail_silently=True)

    return render(request, 'capstone/progress.html', {
        "resource": resource,
        "progress": total_progress,
        "certificate_available": Certificate.objects.filter(user=request.user, resource=resource).exists()
    })

@login_required
@user_passes_test(lambda u: u.user_type == 'mentor' or u.is_superuser)
def upload_chapter_quiz(request):
    """
    Allows mentors to add quizzes to specific chapters of a course.
    """
    resources = Resource.objects.all()
    chapters = Chapter.objects.all()

    if request.method == 'POST':
        resource_id = request.POST.get("resource")
        chapter_id = request.POST.get("chapter")

        resource = get_object_or_404(Resource, id=resource_id)
        chapter = get_object_or_404(Chapter, id=chapter_id)

        quiz_questions = []
        for i in range(1, 6):
            question_text = request.POST.get(f'quiz_question_{i}')
            options = request.POST.get(f'quiz_options_{i}', "").split(',')
            correct_answer = request.POST.get(f'quiz_correct_answer_{i}')

            if question_text and options and correct_answer:
                quiz_questions.append({
                    "question": question_text.strip(),
                    "options": [opt.strip() for opt in options if opt.strip()],
                    "correct_answer": correct_answer.strip()
                })

        if quiz_questions:
            Quiz.objects.create(resource=resource, chapter=chapter, questions=quiz_questions)

        messages.success(request, "Quiz uploaded successfully!")
        return redirect('mentor_dashboard')

    return render(request, 'capstone/upload_chapter_quiz.html', {'resources': resources, 'chapters': chapters})

@login_required
@user_passes_test(lambda u: u.user_type == 'mentor' or u.is_superuser)
def upload_final_quiz(request):
    """Allows mentors to upload a final quiz for a selected course"""
    resources = Resource.objects.all()  # Get all courses for dropdown

    if request.method == 'POST':
        resource_id = request.POST.get('resource')
        resource = get_object_or_404(Resource, id=resource_id)

        quiz_questions = []
        for i in range(1, 11):  # Ensure 10 questions are collected
            question_text = request.POST.get(f'quiz_question_{i}')
            options = request.POST.get(f'quiz_options_{i}', "").split(',')
            correct_answer = request.POST.get(f'quiz_correct_answer_{i}')

            if question_text and options and correct_answer:
                quiz_questions.append({
                    "question": question_text.strip(),
                    "options": [opt.strip() for opt in options if opt.strip()],
                    "correct_answer": correct_answer.strip()
                })

        if quiz_questions:
            # Ensure only ONE final quiz per resource
            final_quiz, created = Quiz.objects.get_or_create(
                resource=resource, 
                is_final_quiz=True, 
                defaults={"questions": quiz_questions}
            )
            if not created:
                final_quiz.questions = quiz_questions
                final_quiz.save()

            messages.success(request, "Final Quiz uploaded successfully!")
            return redirect('mentor_dashboard')

    return render(request, 'capstone/upload_final_quiz.html', {"resources": resources})


@login_required
def reset_quiz(request, quiz_id):
    """Clears previous quiz attempt and redirects back to take the quiz again."""
    quiz = get_object_or_404(Quiz, id=quiz_id)
    progress, created = UserProgress.objects.get_or_create(user=request.user, resource=quiz.chapter.resource)

    # Remove the stored attempt
    if str(quiz_id) in progress.quiz_attempts:
        del progress.quiz_attempts[str(quiz_id)]
        progress.save()

    messages.info(request, "🔄 Your previous attempt has been reset. You can take the quiz again.")

    return redirect('take_quiz', quiz_id=quiz_id)

@login_required
def reset_final_quiz(request, quiz_id):
    """Clears previous attempt for the final quiz and allows retry."""
    quiz = get_object_or_404(Quiz, id=quiz_id, is_final_quiz=True)
    progress, created = UserProgress.objects.get_or_create(user=request.user, resource=quiz.resource)

    # Remove previous attempt for this quiz
    if str(quiz_id) in progress.quiz_attempts:
        del progress.quiz_attempts[str(quiz_id)]
        progress.save()

    messages.info(request, "🔄 Your previous final quiz attempt has been reset. You can take the quiz again.")

    return redirect('take_final_quiz', quiz_id=quiz_id)


@login_required
def download_certificate(request, resource_id):
    certificate = get_object_or_404(Certificate, user=request.user, resource_id=resource_id)
    if certificate.certificate_file:
        return redirect(certificate.certificate_file.url)
    else:
        messages.error(request, "Certificate not found.")
        return redirect('participant_profile', participant_id=request.user.id)


@login_required
def generate_certificate(request, resource_id):
    if request.method == "POST":
        full_name = request.POST.get("full_name")
        course_title = request.POST.get("course_title")

        if not full_name or not course_title:
            messages.error(request, "Please enter your name and select a course.")
            return redirect("course_detail", resource_id=resource_id)

        resource = get_object_or_404(Resource, id=resource_id)

        # ✅ Check if the user already has a certificate
        certificate, created = Certificate.objects.get_or_create(
            user=request.user, resource=resource,
            defaults={"issued_at": now()}
        )

        # ✅ Locate Pacifico font
        pacifico_path = os.path.join(settings.STATICFILES_DIRS[0], "fonts", "Pacifico.ttf")

        # ✅ Ensure Pacifico font file exists
        if not os.path.exists(pacifico_path):
            messages.error(request, "Pacifico font file is missing! Please upload the font to static/fonts.")
            return redirect("course_detail", resource_id=resource_id)

        pdfmetrics.registerFont(TTFont("Pacifico", pacifico_path))

        # ✅ Create directory if it doesn't exist
        certificates_dir = os.path.join(settings.MEDIA_ROOT, "certificates")
        os.makedirs(certificates_dir, exist_ok=True)

        # ✅ Define file path
        pdf_filename = f"certificate_{request.user.username}_{resource.id}.pdf"
        pdf_path = os.path.join(certificates_dir, pdf_filename)

        # ✅ Create certificate in landscape format
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=landscape(letter))
        width, height = landscape(letter)

        # 🎨 **Stylish Border**
        c.setStrokeColor(colors.black)
        c.setLineWidth(5)
        c.rect(25, 25, width - 50, height - 50)

        # 🎖 **Title: Certificate of Completion**
        c.setFont("Helvetica-Bold", 40)
        c.setFillColor(colors.darkblue)
        c.drawCentredString(width / 2, height - 90, "CERTIFICATE OF COMPLETION")
        c.setFillColor(colors.black)

        # 🏆 **Subtitle: This is proudly presented to**
        c.setFont("Helvetica", 18)
        c.drawCentredString(width / 2, height - 150, "This is proudly presented to")

        # 🏅 **Participant's Name in Stylish Font**
        c.setFont("Pacifico", 35)
        c.setFillColor(colors.blue)
        c.drawCentredString(width / 2, height - 200, full_name)
        c.setFillColor(colors.black)

        # 📜 **Course Completion Statement**
        c.setFont("Helvetica", 16)
        c.drawCentredString(width / 2, height - 250, "For successfully completing the course:")

        # 🎓 **Course Title in Bold**
        c.setFont("Helvetica-Bold", 22)
        c.setFillColor(colors.darkred)
        c.drawCentredString(width / 2, height - 280, course_title)
        c.setFillColor(colors.black)

        # 🏛 **Issuer Info**
        c.setFont("Helvetica", 14)
        c.drawCentredString(width / 2, height - 340, "Issued by: Hustle Platform")
        c.drawCentredString(width / 2, height - 360, f"Date: {now().strftime('%Y-%m-%d')}")

        # ✍️ **Signature Line & Placeholder**
        c.line(width / 2 - 120, height - 430, width / 2 + 120, height - 430)
        c.setFont("Pacifico", 50)
        c.setFillColor(colors.blue)
        c.drawCentredString(width / 2, height - 460, "HusP")
        c.setFillColor(colors.black)

        # ✅ Save certificate
        c.showPage()
        c.save()

        buffer.seek(0)
        certificate.certificate_file.save(
            f"certificates/{pdf_filename}", 
            ContentFile(buffer.read())
        )
        certificate.save()

        # ✅ **Send Email Notification with Certificate Link**
        certificate_url = request.build_absolute_uri(certificate.certificate_file.url)
        email_subject = "🎉 Congratulations! Your Certificate is Ready!"
        email_body = f"""
        <p>Dear {request.user.get_full_name()},</p>
        <p>🎉 Congratulations! You have successfully completed the course <strong>{resource.title}</strong>.</p>
        <p>Your certificate is now available for download. Go and Download it from your Profile</p>
        <p>Thank you for learning with Hustle Platform! Keep striving for excellence. 🚀</p>
        <p>Best regards,<br><strong>Hustle Platform Team</strong></p>
        """

        email = EmailMessage(
            subject=email_subject,
            body=email_body,
            from_email=settings.DEFAULT_FROM_EMAIL,
            to=[request.user.email],
        )
        email.content_subtype = "html"
        email.send(fail_silently=True)

        messages.success(request, "🎓 Certificate generated successfully! You can now download it from your profile.")

        # ✅ Return certificate file for download
        return FileResponse(open(pdf_path, "rb"), as_attachment=True, filename=f"{full_name}_certificate.pdf")

    return redirect("course_detail", resource_id=resource_id)

@login_required
@user_passes_test(lambda u: u.user_type == "mentor")
def mentor_add_availability(request):
    if request.method == "POST":
        form = MentorAvailabilityForm(request.POST)
        if form.is_valid():
            availability = form.save(commit=False)
            availability.mentor = request.user
            availability.save()
            
            # ✅ Debugging Statement
            print(f"✅ Availability Saved: {availability.date} {availability.start_time}-{availability.end_time} for {availability.mentor.full_name}")
            
            messages.success(request, "Availability added successfully!")
            return redirect("mentor_dashboard")
        else:
            print("❌ Form errors:", form.errors)  # ✅ Debugging
    else:
        form = MentorAvailabilityForm()

    return render(request, "capstone/mentor_add_availability.html", {"form": form})


def remove_emojis(text):
    emoji_pattern = re.compile(
        "["
        "\U0001F600-\U0001F64F"  # Emoticons
        "\U0001F300-\U0001F5FF"  # Symbols & pictographs
        "\U0001F680-\U0001F6FF"  # Transport & map symbols
        "\U0001F1E0-\U0001F1FF"  # Flags (iOS)
        "]+",
        flags=re.UNICODE,
    )
    return emoji_pattern.sub(r"", text)  # Remove emojis
@login_required
def schedule_session(request, mentor_id, session_id):
    """
    Allows participants to book a session using a mentor's pre-defined availability slot.
    The participant only adds notes; date & time are pre-filled.
    """
    mentor = get_object_or_404(User, id=mentor_id, user_type='mentor')
    available_session = get_object_or_404(MentorAvailability, id=session_id, mentor=mentor)

    # Debugging logs to verify correct data
    print(f"🔍 Mentor: {mentor.full_name}, Session ID: {session_id}")
    print(f"📅 Date: {available_session.date}, Time: {available_session.start_time}, Booked: {available_session.is_booked}")

    if request.method == 'POST':
        form = ScheduleSessionForm(request.POST)

        if form.is_valid():
            session = form.save(commit=False)
            session.participant = request.user
            session.mentor = mentor
            session.date = available_session.date  # ✅ Pre-filled from mentor's slot
            session.start_time = available_session.start_time  # ✅ Pre-filled
            session.end_time = available_session.end_time  # ✅ Pre-filled if needed
            session.notes = form.cleaned_data['notes']
            session.is_confirmed = False
            session.save()

            # Mark session as booked
            available_session.is_booked = True
            available_session.save()

            # ✅ Send notification to mentor
            Notification.objects.create(
                user=mentor,
                message=f"📅 New session request from {request.user.full_name} on {session.date} at {session.start_time}. Notes: {session.notes}"
            )

            # ✅ Send email notification to mentor
            subject = "New Mentorship Session Request"
            message = f"Hello {mentor.first_name},\n\n" \
                      f"You have a new session request from {request.user.full_name}.\n\n" \
                      f"📅 Date: {session.date}\n" \
                      f"⏰ Time: {session.start_time}\n" \
                      f"📝 Notes: {session.notes}\n\n" \
                      f"Please log in to your dashboard to confirm or reschedule this session.\n\n" \
                      f"Best regards,\nYour Mentorship Team"

            send_mail(
                subject,
                message,
                settings.DEFAULT_FROM_EMAIL,
                [mentor.email],
                fail_silently=False,
            )

            messages.success(request, "✅ Session scheduled successfully! The mentor has been notified.")
            return redirect('participant_dashboard')

    else:
        # ✅ Create form with only "notes" (without date & time)
        form = ScheduleSessionForm()

    return render(request, 'capstone/schedule_session.html', {
        'form': form,
        'mentor': mentor,
        'available_session': available_session,
    })


@login_required
def participant_schedule(request):
    """Displays available mentor sessions for participants."""
    available_slots = MentorAvailability.objects.filter(is_booked=False).order_by('date', 'start_time')

    return render(request, 'capstone/participant_schedule.html', {
        'available_slots': available_slots
    })



@login_required
def available_sessions(request):
    """Show available mentor sessions for participants."""
    if request.user.user_type != "participant":
        messages.error(request, "❌ Only participants can access this page.")
        return redirect('participant_dashboard')

    # Show only **available** slots
    available_slots = MentorAvailability.objects.filter(is_booked=False).order_by("date", "start_time")

    return render(request, "capstone/available_sessions.html", {
        "available_slots": available_slots
    })


@login_required
def mentor_sessions(request):
    if request.user.user_type != 'mentor':
        return redirect('mentor_dashboard')

    # ✅ Fetch only unconfirmed sessions
    sessions = ScheduledSession.objects.filter(
        mentor=request.user,
        is_confirmed=False  # Exclude confirmed sessions
    ).select_related('participant').order_by('date', 'start_time')

    print(f"✅ Unconfirmed Sessions Fetched for {request.user.full_name}: {sessions}")  # Debugging

    return render(request, 'capstone/mentor_sessions.html', {'sessions': sessions})


@login_required
@user_passes_test(lambda u: u.user_type == "mentor")
def confirm_session(request, session_id):
    session = get_object_or_404(ScheduledSession, id=session_id, mentor=request.user)

    # ✅ Automatically assign the predefined Google Meet link
    meeting_link = f"https://meet.google.com/hsj-kuen-tky"

    # ✅ Convert `session.start_time` to `datetime.datetime` if it's a `time`
    if isinstance(session.start_time, time):  
        start_time = datetime.combine(session.date, session.start_time)  # Merge date & time
    else:
        start_time = session.start_time  # It's already a full datetime object

    # ✅ Ensure `start_time` is timezone-aware
    if start_time.tzinfo is None:
        start_time = make_aware(start_time)
    else:
        start_time = localtime(start_time)

    # ✅ Define `end_time`
    end_time = start_time + timedelta(hours=1)

    # ✅ Update session in the database
    session.is_confirmed = True
    session.meeting_link = meeting_link  # ✅ Automatically set the meeting link
    session.save()

    # 🔹 Generate Google Calendar Link
    google_calendar_url = (
        "https://www.google.com/calendar/render?"
        + urllib.parse.urlencode({
            "action": "TEMPLATE",
            "text": f"Session with {session.mentor.full_name}",
            "dates": f"{start_time.strftime('%Y%m%dT%H%M%S')}/{end_time.strftime('%Y%m%dT%H%M%S')}",
            "details": f"Organized by Hustle Platform\n\nMeeting with {session.mentor.full_name}\nMeeting Link: {meeting_link}",
            "location": meeting_link,
            "sf": "true",
            "output": "xml",
            "add": f"{settings.DEFAULT_FROM_EMAIL}",
        })
    )

    # 🔹 Generate Outlook Calendar Link
    outlook_calendar_url = (
        "https://outlook.live.com/calendar/0/deeplink/compose?"
        + urllib.parse.urlencode({
            "subject": f"Mentorship Session with {session.mentor.full_name}",
            "startdt": start_time.isoformat(),
            "enddt": end_time.isoformat(),
            "body": f"Organized by Hustle Platform\n\nMeeting with {session.mentor.full_name}\nMeeting Link: {meeting_link}",
            "location": meeting_link,
            "organizer": f"{settings.DEFAULT_FROM_EMAIL}",  # ✅ Correctly format the organizer
        })
    )

    # 🔹 Generate ICS File Content (for Apple & Other Calendar Apps)
    ics_content = f"""BEGIN:VCALENDAR
VERSION:2.0
PRODID:-//Hustle Platform//Mentorship Sessions//EN
BEGIN:VEVENT
ORGANIZER;CN=Hustle Platform:MAILTO:{settings.DEFAULT_FROM_EMAIL}  # ✅ Correct organizer email
SUMMARY:Mentorship Session with {session.mentor.full_name}
DTSTART:{start_time.strftime('%Y%m%dT%H%M%S')}
DTEND:{end_time.strftime('%Y%m%dT%H%M%S')}
DESCRIPTION:Organized by Hustle Platform\\nMeeting with {session.mentor.full_name}\\nMeeting Link: {meeting_link}
LOCATION:{meeting_link}
END:VEVENT
END:VCALENDAR
"""
    ics_filename = f"mentorship_session_{session.id}.ics"

    # 📧 Email Content with Shortened Links
    email_body = f"""
    <p>✅ Your session with <strong>{session.mentor.full_name}</strong> has been confirmed!</p>
    <p><strong>📅 Date:</strong> {session.date}</p>
    <p><strong>⏰ Time:</strong> {start_time.strftime('%H:%M %p')} (CAT)</p>
    <p><strong>🔗 Meeting Link:</strong> <a href="{meeting_link}">Join Meeting</a></p>

    <h3>📅 Add to Calendar:</h3>
    <ul>
        <li><a href="{google_calendar_url}">📅 Add to Google Calendar</a></li>
        <li><a href="{outlook_calendar_url}">📅 Add to Outlook Calendar</a></li>
    </ul>

    <p>Alternatively, you can download the attached <strong>ICS file</strong> to manually add it to your calendar.</p>
    """

    # 📧 Send Email to Participant with Calendar Options
    participant_email = EmailMessage(
        subject="📅 Session Confirmed - Hustle Platform!",
        body=email_body.strip(),
        from_email=settings.DEFAULT_FROM_EMAIL,
        to=[session.participant.email],
    )
    participant_email.attach(ics_filename, ics_content, "text/calendar")
    participant_email.content_subtype = "html"  
    participant_email.send(fail_silently=True)

    # 📧 Send Email to Mentor with Calendar Options
    mentor_email = EmailMessage(
        subject="📅 Your Session is Confirmed - Hustle Platform!",
        body=email_body.strip(),
        from_email=settings.DEFAULT_FROM_EMAIL,
        to=[session.mentor.email],
    )
    mentor_email.attach(ics_filename, ics_content, "text/calendar")
    mentor_email.content_subtype = "html"  
    mentor_email.send(fail_silently=True)

    messages.success(request, "✅ Session confirmed! Participants notified with calendar options.")
    return redirect("mentor_dashboard")


@login_required
def view_mentor_availability(request, mentor_id):
    mentor = get_object_or_404(User, id=mentor_id, user_type='mentor')
    available_slots = MentorAvailability.objects.filter(mentor=mentor)

    return render(request, 'capstone/mentor_availability.html', {
        'mentor': mentor,
        'available_slots': available_slots
    })


@login_required
def schedule_availability(request):
    if request.user.user_type != 'mentor':
        return redirect('mentor_dashboard')

    if request.method == 'POST':
        form = ScheduledSessionForm(request.POST)
        if form.is_valid():
            availability = form.save(commit=False)
            availability.mentor = request.user
            availability.save()
            messages.success(request, "✅ Availability added successfully!")
            return redirect('mentor_dashboard')
    else:
        form = ScheduledSessionForm()

    return render(request, 'capstone/schedule_availability.html', {'form': form})


def extract_text_from_file(file):
    """Extracts text from an uploaded PDF or DOCX file."""
    file_ext = os.path.splitext(file.name)[1].lower()
    
    if file_ext == '.pdf':
        reader = PyPDF2.PdfReader(file)
        text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    elif file_ext == '.docx':
        doc = docx.Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
    else:
        return None  # Unsupported file format
    
    return text


@login_required
def get_quiz(request, quiz_id):
    """Retrieve quiz details including questions and options."""
    quiz = get_object_or_404(Quiz, id=quiz_id)
    questions = quiz.questions.all()

    quiz_data = {
        "quiz_id": quiz.id,
        "title": quiz.title,
        "questions": [
            {
                "id": q.id,
                "question_text": q.question_text,
                "question_type": q.question_type,
                "options": q.options if q.options else []
            }
            for q in questions
        ]
    }
    
    return JsonResponse(quiz_data)

@method_decorator(csrf_exempt, name='dispatch')
@login_required
def submit_quiz_attempt(request, quiz_id):
    """Handle quiz submission and store user responses."""
    if request.method == "POST":
        quiz = get_object_or_404(Quiz, id=quiz_id)
        data = json.loads(request.body)
        user_answers = data.get("answers", {})

        correct_answers = 0
        total_questions = quiz.questions.count()

        for q in quiz.questions.all():
            if q.question_type == "MCQ" and user_answers.get(str(q.id)) == q.correct_answer:
                correct_answers += 1

        score = (correct_answers / total_questions) * 100 if total_questions > 0 else 0

        # Store quiz attempt
        QuizAttempt.objects.create(
            user=request.user,
            quiz=quiz,
            answers=user_answers,
            score=score
        )

        return JsonResponse({"success": True, "score": score})
    
    return JsonResponse({"error": "Invalid request method"}, status=400)


@login_required
@staff_member_required  # Ensures only admins can access
def admin_dashboard(request):
    participants = User.objects.filter(user_type="participant")
    mentors = User.objects.filter(user_type="mentor")
    resources = Resource.objects.all()
    threads = Thread.objects.all()
    job_listings = JobListing.objects.all()
    mentorship_requests = MentorshipRequest.objects.all()
    job_applications = JobApplication.objects.all()
    mentor_availability = MentorAvailability.objects.all()
    feedbacks = Feedback.objects.all()

    # Calculate average ratings for mentors based on feedbacks
    mentor_ratings = Feedback.objects.values('mentor_id').annotate(avg_rating=Avg('rating'))
    
    # Convert to dictionary {mentor_id: avg_rating}
    mentor_ratings_dict = {rating['mentor_id']: rating['avg_rating'] for rating in mentor_ratings}
    


    context = {
        "participants": participants,
        "mentors": mentors,
        "resources": resources,
        "threads": threads,
        "job_listings": job_listings,
        "mentorship_requests": mentorship_requests,
        "job_applications": job_applications,
        "mentor_availability": mentor_availability,
        "feedbacks": feedbacks,
        'mentor_ratings': mentor_ratings_dict,
    }
    return render(request, "capstone/admin_dashboard.html", context)

@login_required
@staff_member_required
def resource_participant_progress(request, resource_id):
    """Displays participant progress for a specific resource."""
    resource = get_object_or_404(Resource, id=resource_id)
    user_progress = UserProgress.objects.filter(resource=resource).select_related("user")

    context = {
        "resource": resource,
        "user_progress": user_progress,
    }
    return render(request, "capstone/resource_participant_progress.html", context)


@login_required
@staff_member_required
def admin_delete_user(request, user_id):
    user = get_object_or_404(User, id=user_id)
    user.delete()
    messages.success(request, "User deleted successfully.")
    return redirect('admin_dashboard')


@login_required
@staff_member_required
def admin_delete_resource(request, resource_id):
    resource = get_object_or_404(Resource, id=resource_id)
    resource.delete()
    messages.success(request, "Resource deleted successfully.")
    return redirect('admin_dashboard')


@login_required
@staff_member_required
def admin_delete_thread(request, thread_id):
    thread = get_object_or_404(Thread, id=thread_id)
    thread.delete()
    messages.success(request, "Thread deleted successfully.")
    return redirect('admin_dashboard')


@login_required
@staff_member_required
def admin_delete_job(request, job_id):
    job = get_object_or_404(JobListing, id=job_id)
    job.delete()
    messages.success(request, "Job listing deleted successfully.")
    return redirect('admin_dashboard')


@login_required
@staff_member_required
def admin_delete_slot(request, slot_id):
    slot = get_object_or_404(MentorAvailability, id=slot_id)
    slot.delete()
    messages.success(request, "Mentor availability slot deleted successfully.")
    return redirect('admin_dashboard')

@login_required
@staff_member_required
def admin_view_resource(request, resource_id):
    resource = get_object_or_404(Resource, id=resource_id)
    chapters = resource.chapters.all()
    quizzes = Quiz.objects.filter(resource=resource)

    return render(request, "capstone/admin_view_resource.html", {
        "resource": resource,
        "chapters": chapters,
        "quizzes": quizzes,
    })

@login_required
@staff_member_required
def admin_view_resource_progress(request, resource_id):
    resource = get_object_or_404(Resource, id=resource_id)
    user_progress = UserProgress.objects.filter(resource=resource)

    participants_progress = []
    for progress in user_progress:
        participants_progress.append({
            "user": progress.user,
            "completed_chapters": progress.completed_chapters.count(),
            "total_chapters": resource.chapters.count(),
            "completed_quizzes": progress.completed_quizzes.count(),
            "total_quizzes": Quiz.objects.filter(resource=resource).count(),
            "final_score": progress.final_score if progress.final_score is not None else "N/A",  # ✅ Final Score Handling
            "progress_percentage": progress.progress_percentage,
            "certificate": Certificate.objects.filter(user=progress.user, resource=resource).exists()  # ✅ Check if user has a certificate
        })

    return render(request, "capstone/admin_view_resource_progress.html", {
        "resource": resource,
        "user_progress": user_progress,
        "participants_progress": participants_progress,
    })


@login_required
@staff_member_required
def admin_view_thread(request, thread_id):
    thread = get_object_or_404(Thread, id=thread_id)
    comments = thread.comments.all()

    return render(request, "capstone/admin_view_thread.html", {
        "thread": thread,
        "comments": comments,
    })


@login_required
@staff_member_required
def admin_delete_comment(request, comment_id):
    """Allows an admin to delete a comment."""
    comment = get_object_or_404(Comment, id=comment_id)
    thread_id = comment.thread.id  # Store thread ID before deletion
    comment.delete()
    
    messages.success(request, "Comment deleted successfully.")
    return redirect("admin_view_thread", thread_id=thread_id)

@login_required
@staff_member_required
def admin_resource_detail(request, resource_id):
    resource = get_object_or_404(Resource, id=resource_id)
    chapters = Chapter.objects.filter(resource=resource)
    final_quiz = Quiz.objects.filter(resource=resource, is_final_quiz=True).first()
    user_progress = UserProgress.objects.filter(resource=resource)

    context = {
        "resource": resource,
        "chapters": chapters,
        "final_quiz": final_quiz,
        "user_progress": user_progress
    }
    return render(request, "capstone/admin_resource_detail.html", context)


@login_required
@staff_member_required
def admin_job_detail(request, job_id):
    """Displays all job applications for a specific job listing."""
    job = get_object_or_404(JobListing, id=job_id)
    applications = JobApplication.objects.filter(job_listing=job)

    context = {
        "job": job,
        "applications": applications,
    }
    return render(request, "capstone/admin_job_detail.html", context)

@login_required
@staff_member_required
def admin_scheduled_session_detail(request, session_id):
    """Displays details of a scheduled session, including participant notes and confirmation status."""
    session = get_object_or_404(ScheduledSession, id=session_id)

    context = {
        "session": session,
    }
    return render(request, "capstone/admin_scheduled_session_detail.html", context)

@login_required
@staff_member_required
def admin_delete_chapter(request, chapter_id):
    chapter = get_object_or_404(Chapter, id=chapter_id)
    resource_id = chapter.resource.id  # Save resource ID before deleting

    # Delete the chapter
    chapter.delete()

    # Redirect back to the resource detail page
    return redirect("admin_resource_detail", resource_id=resource_id)

@login_required
@staff_member_required
def admin_delete_feedback(request, feedback_id):
    feedback = get_object_or_404(Feedback, id=feedback_id)
    mentor_id = feedback.mentor.id  # Save mentor ID before deleting

    # Delete the feedback
    feedback.delete()

    # Redirect back to the mentor's detail page or the admin dashboard
    return redirect("admin_dashboard") 


@login_required
@staff_member_required
def admin_view_user_profile(request, user_id):
    user = get_object_or_404(User, id=user_id)

    # Get certificates (if participant)
    certificates = Certificate.objects.filter(user=user) if user.user_type == "participant" else None

    # ✅ Fetch Enrolled Courses & Progress (if participant)
    enrolled_courses = []
    if user.user_type == "participant":
        user_progress = UserProgress.objects.filter(user=user).select_related("resource")

        for progress in user_progress:
            enrolled_courses.append({
                "resource": progress.resource,
                "progress_percentage": progress.progress_percentage,  # ✅ Track Progress
                "final_score": progress.final_score if progress.final_score is not None else "N/A",  # ✅ Store Final Score
            })

    # Get mentees (if mentor)
    mentees = None
    if user.user_type == "mentor":
        mentees = User.objects.filter(mentorship_requests__mentor=user, mentorship_requests__status="approved").distinct()

    # Get mentor availability slots (if mentor)
    availability_slots = MentorAvailability.objects.filter(mentor=user) if user.user_type == "mentor" else None

    # Get feedback for mentor
    feedbacks = Feedback.objects.filter(mentor=user).select_related("mentee") if user.user_type == "mentor" else Feedback.objects.none()

    # 🔹 Check if there are feedbacks before calculating average rating
    if feedbacks.exists():
        avg_rating = feedbacks.aggregate(Avg("rating"))["rating__avg"]
    else:
        avg_rating = 0  # Default rating if no feedbacks exist

    context = {
        "user": user,
        "certificates": certificates,
        "enrolled_courses": enrolled_courses,
        "mentees": mentees,
        "availability_slots": availability_slots,
        "feedbacks": feedbacks,
        "avg_rating": round(avg_rating, 1),  # Ensuring it's rounded
    }
    return render(request, "capstone/admin_user_profile.html", context)

@login_required
@staff_member_required
def admin_view_final_quiz(request, quiz_id):
    quiz = get_object_or_404(Quiz, id=quiz_id)

    # Ensure questions are correctly retrieved as a QuerySet
    questions = quiz.questions

    context = {
        "quiz": quiz,
        "questions": questions,
    }
    return render(request, "capstone/admin_final_quiz_detail.html", context)


@login_required
@staff_member_required
def admin_delete_final_quiz(request, quiz_id):
    quiz = get_object_or_404(Quiz, id=quiz_id)

    if request.method == "POST":
        quiz.delete()
        messages.success(request, "Final Quiz deleted successfully.")
        return redirect("admin_dashboard")

    return render(request, "capstone/admin_confirm_delete_quiz.html", {"quiz": quiz})


@login_required
@staff_member_required
def admin_resources(request):
    resources = Resource.objects.all()
    return render(request, "capstone/admin_resources.html", {"resources": resources})

@login_required
@staff_member_required
def admin_threads(request):
    threads = Thread.objects.all()
    return render(request, "capstone/admin_threads.html", {"threads": threads})


@login_required
@staff_member_required
def admin_jobs(request):
    jobs = JobListing.objects.all()
    return render(request, "capstone/admin_jobs.html", {"jobs": jobs})


@login_required
@staff_member_required
def admin_mentorship_requests(request):
    mentorship_requests = MentorshipRequest.objects.all()
    return render(request, "capstone/admin_mentorship_requests.html", {"mentorship_requests": mentorship_requests})


@login_required
@staff_member_required
def admin_job_applications(request):
    job_applications = JobApplication.objects.all()
    return render(request, "capstone/admin_job_applications.html", {"job_applications": job_applications})


@login_required
@staff_member_required
def admin_availability(request):
    mentor_availability = MentorAvailability.objects.all()
    return render(request, "capstone/admin_availability.html", {"mentor_availability": mentor_availability})


@login_required
@staff_member_required
def admin_scheduled_sessions(request):
    scheduled_sessions = ScheduledSession.objects.all()
    return render(request, "capstone/admin_scheduled_sessions.html", {"scheduled_sessions": scheduled_sessions})


@login_required
@staff_member_required
def admin_feedbacks(request):
    feedbacks = Feedback.objects.all()
    return render(request, "capstone/admin_feedbacks.html", {"feedbacks": feedbacks})

@login_required
@staff_member_required
def admin_delete_scheduled_session(request, session_id):
    session = get_object_or_404(ScheduledSession, id=session_id)
    session.delete()
    return redirect("admin_scheduled_sessions")

@login_required
@staff_member_required
def admin_delete_mentorship_request(request, request_id):
    mentorship_request = get_object_or_404(MentorshipRequest, id=request_id)
    mentorship_request.delete()
    return redirect("admin_mentorship_requests")  # Ensure this name matches your URL pattern



@login_required
@user_passes_test(lambda u: u.user_type == 'mentor' or u.is_superuser)
def upload_ai_chapter_quiz(request):
    """
    Allows mentors to generate AI-powered quizzes for specific chapters.
    """
    resources = Resource.objects.all()
    chapters = Chapter.objects.all()

    if request.method == 'POST':
        resource_id = request.POST.get("resource")
        chapter_id = request.POST.get("chapter")

        resource = get_object_or_404(Resource, id=resource_id)
        chapter = get_object_or_404(Chapter, id=chapter_id)

        # ✅ Generate AI-based chapter quiz questions
        quiz_questions = generate_quiz_questions(chapter.title, chapter.content, num_questions=5, quiz_type="chapter")

        if quiz_questions:
            quiz, created = Quiz.objects.get_or_create(
                resource=resource,
                chapter=chapter,
                defaults={"questions": quiz_questions}
            )
            if not created:
                quiz.questions = quiz_questions
                quiz.save()

            messages.success(request, "✅ AI-Generated Chapter Quiz uploaded successfully!")
        else:
            messages.error(request, "❌ Failed to generate quiz. Try again.")

        return redirect('mentor_dashboard')

    return render(request, 'capstone/upload_ai_chapter_quiz.html', {'resources': resources, 'chapters': chapters})

@login_required
@user_passes_test(lambda u: u.user_type == 'mentor' or u.is_superuser)
def upload_final_ai_quiz(request):
    """
    Allows mentors to generate AI-powered final quizzes for a full course.
    """
    resources = Resource.objects.all()

    if request.method == 'POST':
        resource_id = request.POST.get("resource")
        resource = get_object_or_404(Resource, id=resource_id)

        # ✅ Generate AI quiz questions for final quiz
        quiz_questions = generate_quiz_questions(resource.title, resource.description, num_questions=10, quiz_type="final")

        if quiz_questions:
            final_quiz, created = Quiz.objects.get_or_create(
                resource=resource,
                is_final_quiz=True,
                defaults={"questions": quiz_questions}
            )
            if not created:
                final_quiz.questions = quiz_questions
                final_quiz.save()

            messages.success(request, "✅ AI-Generated Final Quiz uploaded successfully!")
        else:
            messages.error(request, "❌ Failed to generate final quiz.")

        return redirect('mentor_dashboard')

    return render(request, 'capstone/upload_ai_final_quiz.html', {"resources": resources})
@login_required
@user_passes_test(lambda u: u.user_type == 'mentor' or u.is_superuser)
def ai_generate_chapter_quiz(request, chapter_id):
    """
    API endpoint that generates an AI-based quiz for a chapter.
    """
    chapter = get_object_or_404(Chapter, id=chapter_id)
    quiz_questions = generate_quiz_questions(chapter.title, chapter.content, num_questions=5, quiz_type="chapter")

    if quiz_questions:
        return JsonResponse({"success": True, "questions": quiz_questions})
    else:
        return JsonResponse({"success": False, "error": "Failed to generate quiz."})


@login_required
@user_passes_test(lambda u: u.user_type == 'mentor' or u.is_superuser)
def ai_generate_final_quiz(request, resource_id):
    """
    API endpoint that generates an AI-based final quiz for a course.
    """
    resource = get_object_or_404(Resource, id=resource_id)
    quiz_questions = generate_quiz_questions(resource.title, resource.description, num_questions=10, quiz_type="final")

    if quiz_questions:
        return JsonResponse({"success": True, "questions": quiz_questions})
    else:
        return JsonResponse({"success": False, "error": "Failed to generate quiz."})



@login_required
def mentor_profile_settings(request):
    """Allows mentors to update profile settings including language preference."""
    if request.method == "POST":
        language = request.POST.get("language_preference")
        request.user.language_preference = language
        request.user.save()
        messages.success(request, "Profile settings updated successfully!")
        return redirect("mentor_profile_settings")

    return render(request, "capstone/mentor_profile_settings.html")

@login_required
def participant_profile_settings(request):
    """Allows participants to update profile settings including language preference."""
    if request.method == "POST":
        language = request.POST.get("language_preference")
        request.user.language_preference = language
        request.user.save()
        messages.success(request, "Profile settings updated successfully!")
        return redirect("participant_profile_settings")

    return render(request, "capstone/participant_profile_settings.html")

def set_language_preference(request):
    """Allows both registered and unregistered users to set language preference."""
    if request.method == "POST":
        language = request.POST.get("language_preference")

        if request.user.is_authenticated:
            # ✅ Logged-in user: Save preference to their profile
            request.user.language_preference = language
            request.user.save()
            messages.success(request, "Profile settings updated successfully!")
        else:
            # ✅ Unregistered user: Store preference in session
            request.session["language_preference"] = language
            messages.success(request, "Language preference saved for this session!")

        return redirect("homepage")  # Redirect user to homepage or another page

    return render(request, "capstone/set_language.html")


def create_notification(user, message):
    """Creates a notification and sends an email notification."""
    notification = Notification.objects.create(user=user, message=message)

    # Send an email notification
    send_mail(
        subject="📢 New Notification - Hustle Platform",
        message=f"Hello {user.username},\n\n{message}\n\nBest regards,\nHustle Platform Team",
        from_email=settings.DEFAULT_FROM_EMAIL,
        recipient_list=[user.email],
        fail_silently=False,
    )

def notify_users(message, thread_id):
    """
    Notify all users in their dashboard and via email when a new thread, comment, or vote is created.
    """
    all_users = User.objects.all()

    for user in all_users:
        # ✅ Create Dashboard Notification
        Notification.objects.create(
            user=user,
            message=message,
        )

        # ✅ Send Email Notification
        subject = "🔔 New Activity on Hustle Platform"
        email_message = f"{message}\n\nVisit your dashboard to see what's new🔔🔔🔔!!/"
        send_mail(
            subject,
            strip_tags(email_message),  # Removes HTML
            settings.DEFAULT_FROM_EMAIL,
            [user.email],
            fail_silently=True,
        )